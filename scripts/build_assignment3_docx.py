"""Convert docs/assignment3/assignment3-report.md into a Word document.

Not a full Markdown implementation — handles exactly the constructs used in
the Assignment 3 report: ATX headings, paragraphs, pipe tables, fenced code
blocks, bulleted/numbered lists, images, and inline bold/italic/code/links.

Output: docs/assignment3/Assignment3-Final-Report.docx
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
ASSIGN3_DIR = ROOT / "docs" / "assignment3"
CHARTS = ASSIGN3_DIR / "charts"

# pairs of (markdown-source, docx-target)
DOCUMENTS = [
    (ASSIGN3_DIR / "assignment3-report.md", ASSIGN3_DIR / "Assignment3-Final-Report.docx"),
    (ASSIGN3_DIR / "test-plan.md", ASSIGN3_DIR / "Assignment3-Test-Plan.docx"),
]

# kept for backwards compatibility with the add_image_relative helper
REPORT = DOCUMENTS[0][0]


# ---------- inline formatting ----------

INLINE_TOKEN = re.compile(
    r"(\*\*[^*]+\*\*|"      # **bold**
    r"`[^`]+`|"              # `code`
    r"\*[^*]+\*|"            # *italic* / markdown emphasis
    r"_[^_]+_|"              # _italic_
    r"\[[^\]]+\]\([^)]+\))"  # [text](url)
)


def add_inline(paragraph, text: str) -> None:
    """Append text with simple inline formatting to an existing paragraph."""
    for piece in INLINE_TOKEN.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            run = paragraph.add_run(piece[2:-2])
            run.bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif (piece.startswith("*") and piece.endswith("*")) or (piece.startswith("_") and piece.endswith("_")):
            run = paragraph.add_run(piece[1:-1])
            run.italic = True
        elif piece.startswith("[") and "](" in piece and piece.endswith(")"):
            label, _url = piece[1:-1].split("](", 1)
            run = paragraph.add_run(label)
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            run.underline = True
        else:
            paragraph.add_run(piece)


def strip_md_emphasis(cell: str) -> str:
    """Remove inline markers in table cells to avoid noisy runs — keep text only."""
    cell = cell.strip()
    cell = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell)
    cell = re.sub(r"`([^`]+)`", r"\1", cell)
    cell = re.sub(r"\*([^*]+)\*", r"\1", cell)
    cell = re.sub(r"_([^_]+)_", r"\1", cell)
    cell = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", cell)
    cell = cell.replace("&nbsp;", " ")
    return cell


# ---------- document helpers ----------

def set_base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def add_heading(doc: Document, text: str, level: int) -> None:
    heading = doc.add_heading(text, level=min(level, 4))
    if level == 0 or level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_code_block(doc: Document, lines: list[str]) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.2)
    run = para.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_image_relative(doc: Document, rel_path: str, anchor: Path = None) -> None:
    # Resolve relative to the source markdown file
    base = (anchor or REPORT).parent
    target = (base / rel_path).resolve()
    if not target.exists():
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(target), width=Inches(6.2))


def add_bullet(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="List Bullet")
    add_inline(para, text)


def add_number(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="List Number")
    add_inline(para, text)


def add_paragraph_markdown(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    add_inline(para, text)


# ---------- table parsing ----------

def is_table_separator(line: str) -> bool:
    """Check if a line is the --- separator of a pipe table header."""
    stripped = line.strip()
    if not stripped.startswith("|") and "|" not in stripped:
        return False
    cells = [c.strip() for c in stripped.strip("|").split("|")]
    return all(re.match(r"^:?-+:?$", c) for c in cells) and len(cells) > 0


def parse_pipe_row(line: str) -> list[str]:
    parts = line.strip().strip("|").split("|")
    return [strip_md_emphasis(p) for p in parts]


def add_table_rows(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    n_cols = len(header)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    # Header row
    for idx, cell in enumerate(header):
        p = table.rows[0].cells[idx].paragraphs[0]
        run = p.add_run(cell)
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        # Pad / truncate row to match column count
        padded = (row + [""] * n_cols)[:n_cols]
        for idx, value in enumerate(padded):
            cells[idx].text = value
    doc.add_paragraph("")


# ---------- main parser ----------

def convert(md_path: Path, out_path: Path) -> None:
    doc = Document()
    set_base_style(doc)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Blank line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$|^\*{3,}$|^_{3,}$", stripped):
            i += 1
            continue

        # HTML comment
        if stripped.startswith("<!--"):
            while i < n and "-->" not in lines[i]:
                i += 1
            i += 1
            continue

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            code_lines: list[str] = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            add_code_block(doc, code_lines)
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = strip_md_emphasis(m.group(2))
            add_heading(doc, text, level)
            i += 1
            continue

        # Image
        m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if m:
            add_image_relative(doc, m.group(2), anchor=md_path)
            i += 1
            continue

        # Pipe table
        if stripped.startswith("|") and "|" in stripped and i + 1 < n and is_table_separator(lines[i + 1]):
            header = parse_pipe_row(stripped)
            i += 2  # skip header + separator
            rows: list[list[str]] = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(parse_pipe_row(lines[i]))
                i += 1
            add_table_rows(doc, header, rows)
            continue

        # Bulleted list
        if re.match(r"^[-*]\s+", stripped):
            while i < n and re.match(r"^[-*]\s+", lines[i].strip()):
                add_bullet(doc, lines[i].strip()[2:].strip())
                i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", stripped):
            while i < n and re.match(r"^\d+\.\s+", lines[i].strip()):
                text = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                add_number(doc, text)
                i += 1
            continue

        # Paragraph — collect consecutive non-blank, non-special lines
        buf: list[str] = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].strip()
            if not nxt:
                break
            if re.match(r"^#{1,6}\s", nxt):
                break
            if nxt.startswith("```") or nxt.startswith("|") or nxt.startswith("!["):
                break
            if re.match(r"^[-*]\s+", nxt) or re.match(r"^\d+\.\s+", nxt):
                break
            buf.append(nxt)
            i += 1
        add_paragraph_markdown(doc, " ".join(buf))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    for src, dst in DOCUMENTS:
        if not src.exists():
            print(f"skip: {src} not found")
            continue
        convert(src, dst)
        print(f"wrote {dst}")


if __name__ == "__main__":
    main()
